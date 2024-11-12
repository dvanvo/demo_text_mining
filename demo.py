import streamlit as st
from keybert import KeyBERT
from docx import Document
from sklearn.feature_extraction.text import CountVectorizer,TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import itertools
import os
from io import BytesIO
from collections import defaultdict, Counter
import re
import pandas as pd
import py_vncorenlp  # for VnCoreNLP vietnamese processing
import math
import networkx as nx  # Lib for graph
from itertools import combinations
import matplotlib.pyplot as plt
import numpy as np
from wordcloud import WordCloud
from matplotlib import use  
from sklearn.cluster import KMeans
import evaluate
import altair as alt

# Streamlit start
st.set_page_config(page_title="Summary & Keywords Extraction", layout="wide")



""" Introduction:

    This is programe demo for text mining.
    
    Project
    -	Objectif: building a module for keywords extraction
    -	Input: individual document
    -	Output: keywords, Summarize text
    -	Implementation plan
        1. Choosing TextRank a graph-based method to implement.
        2. Implementing a KeyBERT method (using cosine similarity on BERT embedding method).
        3. Analyzing the obtained results.
        4. If upload multi-documents, you can setup number of clusters
        5. Extraction keywords and summarize text for clusters.

    """
@st.cache_resource  # Ensures the model loads only once
def load_vncorenlp_model():
    vncorenlp_dir = r"D:/CodePy/demo_text_mining/vncorenlp/"
    model = py_vncorenlp.VnCoreNLP(save_dir=vncorenlp_dir, annotators=["wseg"])
    return model

model = load_vncorenlp_model()
# Nhập đường dẫn tới thư mục chứa các tệp .docx
folder_path = r"D:/CodePy/txtmining/1vitinh/"

# Stopwords Vietnamese
stopwords_file  = r"D:/CodePy/demo_text_mining/vietnamese-stopwords.txt"
 

try:
    with open(stopwords_file, 'r', encoding='utf-8') as file:
        STOPWORDS = file.read().splitlines()
except FileNotFoundError:
    st.error("File stopwords does not exist.")
    STOPWORDS = []

#Show các file trong thư mục
def list_file(folder_path):
    print(f"Danh sách các file: ")
    for file_name in os.listdir(folder_path):
        print(f"           {file_name}")

list_file(folder_path)

#Read file docx
def read_docx_file_path(file_path): 
    doc = Document(file_path)
    text = " ".join([para.text for para in doc.paragraphs])
    return text

def read_docx_file_upload(uploaded_file):
    byte_data = io.BytesIO(uploaded_file.read())
    doc = Document(byte_data)
    text = " ".join([para.text for para in doc.paragraphs])
    return text


# Word segment VnCoreNLP
def tokenize_and_filter(text): 
    
    tokens = model.word_segment(text)  # return words tokens list

    filtered_text = []
    for sentence in tokens:
        if isinstance(sentence, str):  # If `sentence` is string
            sentence = sentence.split()  # Segment (split) sentence to tokens

        filtered_sentence = [word for word in sentence if word.lower() not in STOPWORDS] 
        filtered_text.extend(filtered_sentence)  
    
    return filtered_text  # Return list tokens

#Built graph, link tokens together

     
def build_graph(tokens, window_size=2):
    
    graph = nx.Graph()  # init

    # Create edges between words within the context window range
    for i in range(len(tokens) - window_size + 1):
        window = tokens[i:i + window_size]
        for word1, word2 in combinations(window, 2):
            if graph.has_edge(word1, word2):
                graph[word1][word2]['weight'] += 1
            else:
                graph.add_edge(word1, word2, weight=1)

    return graph

def textrank(graph, max_iter=100, damping=0.85):
    
    return nx.pagerank(graph, max_iter=max_iter, alpha=damping)

# Extract keywords textrank:

#     - Extract keywords based on threshold selection strategy.
#     - Mothod: 'relative' or 'mean_std'
   
def extract_keywords_textrank(text, strategy='relative', top_percent=0.1):
    tokens = tokenize_and_filter(text)

    graph = build_graph(tokens)  
    scores = textrank(graph)  
 
    top_n = max(1, int(len(tokens) * top_percent))
    
    if strategy == 'relative':
        # Selcct top `top_percent` keywords by TextRank score
        sorted_scores = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        #top_n = int(len(sorted_scores) * top_percent)  # Number of keywords
        filtered_keywords = sorted_scores[:top_n]

    elif strategy == 'mean_std':
        # Average and standard deviation
        mean_score = np.mean(list(scores.values()))
        std_score = np.std(list(scores.values()))

        # Filter words with a score of >= average + 1 standard deviation
        filtered_keywords = {word: score for word, score in scores.items() 
                             if score >= mean_score + std_score}
        # Sort keywords by descending point
        filtered_keywords = sorted(filtered_keywords.items(), key=lambda x: x[1], reverse=True)
 

    return filtered_keywords

#Generate a title from the text using top keywords

def generate_title(text, keywords, max_words=10):
    
    # Get top keywords for title
    top_keywords = [kw[0] for kw in keywords[:max_words]]
    
    # Create title by capitalizing and joining keywords
    title = " ".join(word.capitalize() for word in top_keywords)
    
    return title

    # summarize_text:

    # Summarize the text by selecting the sentences with the most keywords.
    # - text: The text to be summarized.
    # - keywords: List of keywords [(word, weight)].
    # - top_n: Number of sentences you want to keep in the summary.
   

# def summarize_text(text, keywords, top_n=3):
    
#     # Split sentences with spaCy  
    
#     sentences = re.split(r'(?<=[.!?])\s+', text)  

#     # Create a weighted dictionary of keywords
#     keyword_weights = {kw[0].lower(): kw[1] for kw in keywords}

#     # Cal score for each sentence
#     sentence_scores = {}
#     for sentence in sentences:
#         words = sentence.lower().split()
#         score = sum(keyword_weights.get(word, 0) for word in words)
#         sentence_scores[sentence] = score

#     # Select `top_n` the sentence with the highest score
#     top_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:top_n]
#     summary = " ".join([sent for sent, score in top_sentences])

#     return summary



# Summarize text with title generation
def summarize_text(text, keywords, top_n=3):
    # Generate title  
    title = generate_title(text, keywords)
    
    # Split sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    # Create weighted dictionary of keywords
    keyword_weights = {kw[0].lower(): kw[1] for kw in keywords}
    
    # Calculate score for each sentence with its original position
    sentence_scores = []
    for idx, sentence in enumerate(sentences):
        words = sentence.lower().split()
        score = sum(keyword_weights.get(word, 0) for word in words)
        sentence_scores.append((sentence, score, idx))
    
    # Select top_n sentences based on scores
    top_sentences = sorted(sentence_scores, key=lambda x: x[1], reverse=True)[:top_n]
    
    # Sort selected sentences by their original position
    summary_sentences = sorted(top_sentences, key=lambda x: x[2])
    
    # Join sentences in original order
    summary = " ".join([sent for sent, score, idx in summary_sentences])
    
    return {
        "title": title,
        "summary": summary
    }
#Graph function with vertices as words and display the PageRank score on each vertex. 
def plot_graph_with_scores(graph, scores):
    
    plt.figure(figsize=(12, 8))  

    pos = nx.spring_layout(graph, seed=42) 

    # Draw edges with weights
    nx.draw_networkx_edges(graph, pos, alpha=0.5)

    # Draw nodes (vertices) with a size that depends on the TextRank point
    node_size = [v * 5000 for v in scores.values()]  
    nx.draw_networkx_nodes(graph, pos, node_size=node_size, node_color='lightblue')

    # Draw labels for vertices (keywords) and display TextRank points
    nx.draw_networkx_labels(
        graph, pos, 
        labels={node: f"{node}\n{scores[node]:.4f}" for node in graph.nodes()}, 
        font_size=10, font_color='black'
    )

    plt.title("TextRank Graph - Visualization", fontsize=16)
    plt.axis("off")  
    #plt.show()
    st.pyplot(plt)

#
#............Xu ly cung luc nhieu tai lieu--------------
#

# Đọc tất cả các tài liệu từ folder_path
def read_all_documents(folder_path):
    documents = {}
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            file_path = os.path.join(folder_path, file_name)
            text = read_docx_file_path(file_path)
            documents[file_name] = text
    return documents

# Phân cụm các tài liệu
def cluster_documents(documents, n_clusters=3, min_df=2, max_df=0.95):
    #vectorizer = TfidfVectorizer(stop_words=STOPWORDS)
    vectorizer = TfidfVectorizer(
        stop_words=STOPWORDS,
        min_df=min_df,  # Bỏ qua các từ xuất hiện ít hơn min_df documents
        max_df=max_df,  # Bỏ qua các từ xuất hiện trong hơn max_df% documents
        lowercase=True,
        norm='l2'
    )
    #doc_vectors = vectorizer.fit_transform(documents.values())
    try:
        doc_vectors = vectorizer.fit_transform(documents.values())
    except Exception as e:
        raise ValueError(f"Lỗi khi vectorize documents: {str(e)}")
    
    n_clusters = min(n_clusters, len(documents))
    # K-means clustering
    # Thực hiện K-means clustering
    kmeans = KMeans(
        n_clusters=n_clusters,
        random_state=42,
        n_init=10  # Số lần chạy với các điểm khởi tạo khác nhau
    )

    #kmeans = KMeans(n_clusters=n_clusters, random_state=42)
    clusters = kmeans.fit_predict(doc_vectors)
    
    # Gắn nhãn cụm vào tài liệu
    clustered_docs = {}
    for doc_id, cluster in zip(documents.keys(), clusters):
        if cluster not in clustered_docs:
            clustered_docs[cluster] = []
        clustered_docs[cluster].append((doc_id, documents[doc_id]))
    
    return clustered_docs

#trich keyword nhieu tai lieu
def extract_keywords_multiple_documents(documents, strategy='relative', top_percent=0.1):
    keywords_by_document = {}
    for doc_name, text in documents.items():
        keywords = extract_keywords_textrank(text, strategy=strategy, top_percent=top_percent)
        keywords_by_document[doc_name] = keywords
    return keywords_by_document

#Hiển thị keyword theo tài liệu
def display_keywords(keywords_by_document):
    for doc_name, keywords in keywords_by_document.items():
        st.write(f"**Keywords for {doc_name}:**")
        st.write([kw[0] for kw in keywords])

#Tóm tắt nhiều tài liệu theo cụm
def summarize_documents_by_cluster(clustered_docs, top_n=3):
    summaries_by_cluster = {}
    for cluster_id, docs in clustered_docs.items():
        summaries = []
        for doc_name, text in docs:
            keywords = extract_keywords_textrank(text)
            summary = summarize_text(text, keywords, top_n=top_n)
            summaries.append((doc_name, summary))
        summaries_by_cluster[cluster_id] = summaries
    return summaries_by_cluster

# Hiển thị tóm tắt theo cụm
def display_summaries(summaries_by_cluster):
    for cluster_id, summaries in summaries_by_cluster.items():
        st.write(f"**Cluster {cluster_id} Summaries:**")
        for doc_name, summary in summaries:
            st.write(f"**{doc_name}:** {summary}")
#graph = build_graph(tokens)   
#scores = textrank(graph) 
#plot_graph_with_scores(graph, scores)


# model_name: 
#         - all-MiniLM-L6-v2:
#             Accuracy: Medium
#             Speed: Fast
#             Size: Small
#             Usage: Real-time extraction, large datasets
#         - distiluse-base-multilingual-cased-v2:
#             Accuracy: Medium
#             Speed:  Fast
#             Size: Small
#             Usage: Multilingual, non-English texts
#         - paraphrase-mpnet-base-v2:
#             Accuracy: High
#             Speed: Moderate
#             Size: Medium
#             Usage: Balanced extraction with good performance
#         - paraphrase-distilroberta-base-v1:
#             Accuracy: High
#             Speed: Fast
#             Size: Medium
#             Usage: Speed-optimized with good accuracy
#         - all-mpnet-base-v2: 
#             Accuracy: Very High
#             Speed: Slow
#             Size: Large
#             Usage: esearch and high-quality extraction
#         - sentence-t5-base:
#             Accuracy: Very High
#             Speed: Slow
#             Size: Large
#             Usage: NLP research and top-tier quality


def get_keyword_model(model_name="paraphrase-mpnet-base-v2"): #all-MiniLM-L6-v2
    return KeyBERT(model_name)
 
# init KeyBERT model
#kw_model = get_keyword_model("paraphrase-distilroberta-base-v1")


# Join tokens into cleaned text for keyword extraction
def preprocess_for_keybert(tokens):
    return " ".join(tokens)  # Convert token list to string

# Visualize keywords using WordCloud
def visualize_keywords(keywords):
    word_freq = {kw: score for kw, score in keywords}
    wordcloud = WordCloud(font_path='arial', background_color='white').generate_from_frequencies(word_freq)

    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    #plt.show()
    st.pyplot(plt)

def extract_keywords_keybert(model,text,top_percent=0.1):
    
    tokens = tokenize_and_filter(text)
    cleaned_text = preprocess_for_keybert(tokens)

    # Initialize KeyBERT model 
    kw_model = get_keyword_model(model)
    # Calculate the number of keywords to extract (e.g., 10% of tokens)
    #keyword_percentage = top_percent  # 10%
    num_keywords = max(1, int(len(tokens) * top_percent))  # Ensure at least 1 keyword
     

    # Extract keywords with MMR for diversity
    keywords_mmr = kw_model.extract_keywords(
        cleaned_text,
        keyphrase_ngram_range=(1, 2),  # Extract unigrams and bigrams
        stop_words=list(STOPWORDS),
        use_mmr=True,  # Maximal Marginal Relevance for relevance and diversity
        diversity=0.7,  # Adjust diversity between 0 and 1
        top_n=num_keywords  # Extract keywords based on percentage
    )

    return keywords_mmr

# Trích xuất từ khóa từ nhiều tài liệu sử dụng KeyBERT
def extract_keywords_keybert_multiple_documents(documents, model="paraphrase-mpnet-base-v2", top_percent=0.1):
    keywords_by_document = {}
    for doc_name, text in documents.items():
        keywords = extract_keywords_keybert(model, text, top_percent=top_percent)
        keywords_by_document[doc_name] = keywords
    return keywords_by_document

# Hiển thị từ khóa từ từng tài liệu
def display_keywords_keybert(keywords_by_document):
    for doc_name, keywords in keywords_by_document.items():
        st.write(f"**Keywords for {doc_name}:**")
        st.write([kw[0] for kw in keywords])

def extract_keywords_by_cluster(clustered_docs, model="paraphrase-mpnet-base-v2", top_percent=0.1):
    keywords_by_cluster = {}
    for cluster_id, docs in clustered_docs.items():
        combined_text = " ".join([text for _, text in docs])
        keywords = extract_keywords_keybert(model, combined_text, top_percent=top_percent)
        keywords_by_cluster[cluster_id] = keywords
    return keywords_by_cluster

# Hiển thị từ khóa theo từng cụm
def display_keywords_by_cluster(keywords_by_cluster):
    for cluster_id, keywords in keywords_by_cluster.items():
        st.write(f"**Keywords for Cluster {cluster_id}:**")
        st.write([kw[0] for kw in keywords])
# Tóm tắt nhiều tài liệu trong cụm dựa trên từ khóa KeyBERT
# def summarize_documents_by_cluster_keybert(clustered_docs, model="all-MiniLM-L6-v2", top_percent=0.1, top_n=3):
#     summaries_by_cluster = {}
#     for cluster_id, docs in clustered_docs.items():
#         combined_text = " ".join([text for _, text in docs])
#         keywords = extract_keywords_keybert(model, combined_text, top_percent=top_percent)
        
#         summary = summarize_text(combined_text, keywords, top_n=top_n)
#         summaries_by_cluster[cluster_id] = summary
#     return summaries_by_cluster

#Summarize documents by cluster with titles
def summarize_documents_by_cluster_keybert(clustered_docs, model="paraphrase-mpnet-base-v2", top_percent=0.1, top_n=3):
     
    summaries_by_cluster = {}
    for cluster_id, docs in clustered_docs.items():
        cluster_summaries = []
        
        # Process each document in cluster
        for doc_name, text in docs:
            keywords = extract_keywords_keybert(model, text, top_percent=top_percent)
            summary_data = summarize_text(text, keywords, top_n=top_n)
            cluster_summaries.append({
                "doc_name": doc_name,
                "title": summary_data["title"],
                "summary": summary_data["summary"]
            })
            
        # Generate cluster-level summary
        combined_text = " ".join([text for _, text in docs])
        cluster_keywords = extract_keywords_keybert(model, combined_text, top_percent=top_percent)
        cluster_summary = summarize_text(combined_text, cluster_keywords, top_n=top_n)
        
        summaries_by_cluster[cluster_id] = {
            "cluster_title": cluster_summary["title"],
            "cluster_summary": cluster_summary["summary"],
            "documents": cluster_summaries
        }
    
    return summaries_by_cluster

# Hiển thị tóm tắt theo cụm
# def display_summaries_by_cluster(summaries_by_cluster):
#     for cluster_id, summary in summaries_by_cluster.items():
#         st.write(f"**Summary for Cluster {cluster_id}:**")
#         st.write(summary)

#Display summaries with titles by cluster
def display_summaries_by_cluster(summaries_by_cluster):
    
    for cluster_id, cluster_data in summaries_by_cluster.items():
        st.markdown(f"### 📑 Cluster {cluster_id + 1}: {cluster_data['cluster_title']}")
        
        # Display cluster summary
        st.markdown("#### 📌 Cluster Overview:")
        st.write(cluster_data['cluster_summary'])
        
        # Display individual document summaries
        st.markdown("#### 📄 Document Summaries:")
        for doc_summary in cluster_data['documents']:
            with st.expander(f"{doc_summary['doc_name']} - {doc_summary['title']}"):
                st.write(doc_summary['summary'])

#Hiển thị kết quả đánh giá trong Streamlit
def display_evaluation_results(keyword_df: pd.DataFrame, summary_df: pd.DataFrame):
    
    st.header("📊 Evaluate result")
    
    # Hiển thị đánh giá từ khóa
    st.subheader("🔑 Evaluate keywords extraction")
    #st.dataframe(keyword_df.style.format("{:.4f}"))
    keyword_df_formatted = keyword_df.applymap(lambda x: f"{x:.4f}" if isinstance(x, (int, float)) else x)
    st.dataframe(keyword_df_formatted)
    
    # Tạo biểu đồ so sánh
    keyword_chart_data = pd.melt(keyword_df, id_vars=['Metric'], var_name='Method', value_name='Score')
    keyword_chart = alt.Chart(keyword_chart_data).mark_bar().encode(
        x='Method',
        y='Score',
        color='Method',
        column='Metric'
    ).properties(width=150)
    st.altair_chart(keyword_chart)
    
    # Hiển thị đánh giá tóm tắt
    st.subheader("📝 Evaluate summarize text")
    #st.dataframe(summary_df.style.format("{:.4f}"))
    summary_df_formatted = summary_df.applymap(lambda x: f"{x:.4f}" if isinstance(x, (int, float)) else x)
    st.dataframe(summary_df_formatted)

    # Tạo biểu đồ so sánh cho tóm tắt
    summary_chart_data = pd.melt(summary_df, id_vars=['Metric'], var_name='Method', value_name='Score')
    summary_chart = alt.Chart(summary_chart_data).mark_bar().encode(
        x='Method',
        y='Score',
        color='Method',
        column='Metric'
    ).properties(width=100)
    st.altair_chart(summary_chart)

def read_docx_file(uploaded_file):
    try:
        # Đọc toàn bộ nội dung của file
        file_content = uploaded_file.read()
        
        # Tạo BytesIO object từ nội dung file
        byte_stream = BytesIO(file_content)
        
        # Tạo Document object từ BytesIO stream
        doc = Document(byte_stream)
        
        # Trả file pointer về đầu để có thể đọc lại nếu cần
        uploaded_file.seek(0)
        
        # Extract text từ tất cả paragraphs
        text = " ".join([para.text for para in doc.paragraphs])
        
        return text
        
    except Exception as e:
        # Reset file pointer
        uploaded_file.seek(0)
        raise Exception(f"Error processing document: {str(e)}")

# Sidebar - select document mode and options
st.sidebar.title("🔍 Options")
doc_mode = st.sidebar.radio("Select Document Mode:", ["Single Document", "Multiple Documents"])
top_percent = st.sidebar.slider("% of keywords:", 0.05, 0.5, 0.1)
num_sentences = st.sidebar.slider("Number of summary sentences:", 1, 10, 3)
num_clusters = st.sidebar.slider("Number of clusters:", 1, 10, 1)

# Main Header
st.markdown("<h1 style='text-align: center; color: #4CAF50;'>Demo Application for Extracting Keywords and Summarizing Text</h1>", unsafe_allow_html=True)

# File upload
uploaded_files = st.file_uploader("📂 Upload DOCX file(s)", type=["docx"], help="Only support DOCX files.", accept_multiple_files=(doc_mode == "Multiple Documents"))

if uploaded_files:
    documents = {}
    
    # Xử lý trường hợp single document
    if isinstance(uploaded_files, st.runtime.uploaded_file_manager.UploadedFile):
        uploaded_files = [uploaded_files]
        
    for uploaded_file in uploaded_files:
        try:
            # Kiểm tra file type
            if not uploaded_file.name.endswith('.docx'):
                st.error(f"⚠️ File '{uploaded_file.name}' does'nt file DOCX.")
                continue
                
            # Đọc nội dung file
            doc_text = read_docx_file(uploaded_file)
            
            # Lưu vào dictionary
            doc_name = uploaded_file.name
            documents[doc_name] = doc_text
            
        except Exception as e:
            st.error(f"⚠️ Error read file '{uploaded_file.name}': {str(e)}")
            continue

    # Hiển thị nội dung dựa vào chế độ single/multiple
    if len(documents) == 1:
        # Single document - hiển thị trực tiếp 
        doc_name, doc_text = next(iter(documents.items()))
        st.markdown(f"### 📄 The content  {doc_name}")
        with st.expander("View file's contents", expanded=True):
            st.write(doc_text)
        
        if st.button("🚀 Extract Keywords and Summarize Text"): 
            # Xử lý trích xuất từ khóa và tóm tắt cho một tài liệu
            with st.spinner("Processing..."):
                # Trích xuất từ khóa sử dụng cả TextRank và KeyBERT
                st.markdown("### 🔑 The keywords of document")
                
                # TextRank keywords
                with st.expander("TextRank Keywords", expanded=True):
                    textrank_keywords = extract_keywords_textrank(doc_text, strategy='relative', top_percent=top_percent)
                    st.markdown("\n".join([f"- **{kw[0]}**: {kw[1]:.4f}" 
                                        for kw in textrank_keywords]))
                
                # KeyBERT keywords
                with st.expander("KeyBERT Keywords", expanded=True):
                    keybert_keywords = extract_keywords_keybert("paraphrase-mpnet-base-v2", doc_text, top_percent=top_percent)
                    st.markdown("\n".join([f"- **{kw[0]}**: {kw[1]:.4f}" 
                                        for kw in keybert_keywords]))
                
                #Summarize with TextRank and KeyBERT
                st.subheader("TextRank and KeyBERT Summaries")
                summary_textrank = summarize_text(
                    doc_text,
                    textrank_keywords,
                    top_n=num_sentences
                )
                st.markdown("**TextRank Summary:**")
                with st.expander("View the summarize text", expanded=True):
                    if summary_textrank.get("title"):
                        st.markdown(f"**The title:** {summary_textrank['title']}")
                    st.write(summary_textrank["summary"]) 


                summary_keybert = summarize_text(
                    doc_text,
                    keybert_keywords,
                    top_n=num_sentences
                )

                st.markdown("**KeyBERT Summary:**")
                with st.expander("View the summarize text", expanded=True):
                    if summary_keybert.get("title"):
                        st.markdown(f"**The title:** {summary_keybert['title']}")
                    st.write(summary_keybert["summary"]) 

                # Tóm tắt văn bản
                st.markdown("### 📝 The summarize text - Combine keywords from both methods")
                
                # Combine keywords from both methods for better summary
                combined_keywords = list(set([kw[0] for kw in textrank_keywords + keybert_keywords]))
                weighted_keywords = [(kw, 1.0) for kw in combined_keywords]  # Using equal weights
                
                summary = summarize_text(
                    doc_text,
                    weighted_keywords,
                    top_n=num_sentences
                )
                
                with st.expander("View the summarize text", expanded=True):
                    if summary.get("title"):
                        st.markdown(f"**The title:** {summary['title']}")
                    st.write(summary["summary"]) 
        if st.button("🚀 Graph TextRank and visualiz keywords keyBERT"): 
            with st.spinner("Processing..."):
                
                # Built and show TextRank graph
                tokens = tokenize_and_filter(doc_text)
                graph = build_graph(tokens)
                scores = textrank(graph)

                st.markdown("### 📈 Graph TextRank")
                plot_graph_with_scores(graph, scores)

                # Built and show visualiz keywords keyBERT     
                keywords_keybert= extract_keywords_keybert("paraphrase-mpnet-base-v2", doc_text, top_percent=top_percent) 
                st.markdown("### 📈 visualiz keywords keyBERT")    
                visualize_keywords(keywords_keybert)

        if st.button("🚀 Evalueate Textrank and keyBERT"): 
            with st.spinner("Processing..."):
                keywords_textrank = extract_keywords_textrank(doc_text, strategy='relative', top_percent=top_percent)
                keywords_keybert= extract_keywords_keybert("paraphrase-mpnet-base-v2", doc_text, top_percent=top_percent)
                
                summary_textrank = summarize_text(
                    doc_text,
                    keywords_textrank,
                    top_n=num_sentences
                )
                summary_keybert = summarize_text(
                    doc_text,
                    keywords_keybert,
                    top_n=num_sentences
                )

                textrank_keywords = [keyword for keyword, score in keywords_textrank]
                keybert_keywords = [keyword for keyword, score in keywords_keybert]
                #textrank_keywords = [kw[0] for kw in keywords_textrank]
                #keybert_keywords = [kw[0] for kw in keywords_keybert]


                text = doc_text
                textrank_summary = summary_textrank["summary"]
                keybert_summary = summary_keybert["summary"]

                # Tạo đối tượng evaluator
                evaluator = evaluate.KeywordSummaryEvaluator()
                

                # Tạo báo cáo đánh giá
                keyword_df, summary_df = evaluator.generate_evaluation_report(
                    text, 
                    textrank_keywords, 
                    keybert_keywords, 
                    textrank_summary, 
                    keybert_summary
                )
                # st.write(textrank_keywords)
                # st.write(keybert_keywords)
                # st.write(keyword_df)
                # st.write(summary_df) 


                # Hiển thị kết quả đánh giá
                display_evaluation_results(keyword_df, summary_df)

    else:
        # Multiple documents - hiển thị danh sách với links
        st.markdown("### 📂 List documents upload")
        
        # Phân cụm tài liệu
        n_clusters = min(num_clusters, len(documents))  # Số cụm tối đa là 3 hoặc số lượng tài liệu
        clustered_docs = cluster_documents(documents, n_clusters=n_clusters)
        
        # Hiển thị tài liệu theo cụm
        for cluster_id, docs in clustered_docs.items():
            st.markdown(f"#### 📑 Cluster {cluster_id + 1}")
            for doc_name, text in docs:
                doc_link = f"[Open {doc_name} in new page](?file={doc_name})"
                st.markdown(f"- **{doc_name}** - {doc_link}")

        # Kiểm tra URL parameter để mở tài liệu trong tab mới
        if "file" in st.query_params:
            selected_file = st.query_params["file"][0]
            if selected_file in documents:
                st.markdown(f"### 📄The content of {selected_file}")
                st.write(documents[selected_file])

        # Nút để xử lý trích xuất từ khóa và tóm tắt
        if st.button("🚀 Extract Keywords and Summarize Text"):
            with st.spinner("Processing..."):
                 
                # 2. Trích xuất và hiển thị từ khóa theo cụm
                st.markdown("### 🔍 Cluster keywords")
                keywords_by_cluster = extract_keywords_by_cluster(clustered_docs, 
                                                              model="paraphrase-mpnet-base-v2", 
                                                              top_percent=top_percent)
                for cluster_id, keywords in keywords_by_cluster.items():
                    with st.expander(f"Cluster keywords {cluster_id + 1}"):
                        st.markdown("\n".join([f"- **{kw[0]}**: {kw[1]:.4f}" for kw in keywords]))

                # 3. Tóm tắt văn bản theo cụm
                st.markdown("### 📝 Cluster Summarize")
                summaries = summarize_documents_by_cluster_keybert(clustered_docs, 
                                                               model="paraphrase-mpnet-base-v2", 
                                                               top_percent=top_percent, 
                                                               top_n=num_sentences)
                
                display_summaries_by_cluster(summaries)
 
                

else:
    st.info("Upload file DOCX to starting.")

 